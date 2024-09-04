from docx import Document

def create_worksheet():
    # Erstellen eines neuen Dokuments
    doc = Document()
    
    # Titel hinzufügen
    doc.add_heading('Systemadministration: Problemlösungsfragen', level=1)
    
    # Einleitung hinzufügen
    doc.add_paragraph(
        "In diesem Arbeitsblatt sollen Sie systemische Fragen formulieren, um ein Netzwerkdruckerproblem zu lösen. "
        "Der Drucker reagiert nicht auf Druckaufträge, und das IT-Team ist sich über die Ursache des Problems unsicher. "
        "Verwenden Sie die folgenden Fragen als Leitfaden, um das Problem zu analysieren und mögliche Lösungen zu finden."
    )
    
    # Fragen hinzufügen
    questions = {
        "Zirkuläre Fragen": [
            "Wie hat sich das Verhalten des Druckers in den letzten Wochen verändert, und gibt es eine mögliche Verbindung zwischen diesen Veränderungen und dem aktuellen Problem?"
        ],
        "Lösungsorientierte Fragen": [
            "Was haben wir bisher getan, um das Problem zu lösen, und welche Maßnahmen haben sich als hilfreich herausgestellt?"
        ],
        "Hypothetische Fragen": [
            "Angenommen, der Drucker würde wieder funktionieren, welche Schritte sollten wir unternehmen, um sicherzustellen, dass das Problem nicht erneut auftritt?"
        ],
        "Wunderfragen": [
            "Wenn über Nacht ein Wunder geschehen würde und der Drucker wieder ohne Probleme funktionieren würde, was wäre die erste Veränderung, die wir bemerken würden?"
        ],
        "Begründungsfragen": [
            "Warum könnten die Netzwerkverbindungen des Druckers gestört sein? Welche möglichen Ursachen könnten wir in Betracht ziehen?"
        ],
        "Skalierungsfragen": [
            "Auf einer Skala von 1 bis 10, wie stark schätzen wir das Risiko ein, dass das Problem bei zukünftigen Druckaufträgen wieder auftritt?"
        ]
    }
    
    # Fragen in das Dokument einfügen
    for category, qs in questions.items():
        doc.add_heading(category, level=2)
        for q in qs:
            doc.add_paragraph(q)
    
    # Dokument speichern
    doc.save('Problemlösungsfragen_Arbeitsblatt.docx')

# Skript ausführen
create_worksheet()
